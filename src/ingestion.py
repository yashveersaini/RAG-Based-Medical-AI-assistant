import os
import tempfile
from pathlib import Path
from typing import List
from dataclasses import dataclass
from dotenv import load_dotenv
from langchain_core.documents import Document
from langchain_community.document_loaders import PyMuPDFLoader, TextLoader
from langchain_classic.text_splitter import RecursiveCharacterTextSplitter
from langchain_huggingface import HuggingFaceEmbeddings
from pinecone import Pinecone
from langchain_pinecone import PineconeVectorStore
from pinecone import ServerlessSpec
from docx import Document as DocxDocument
from config import settings


# ── Supported formats ────────────────────────────────────────────
SUPPORTED_EXTENSIONS = {".pdf", ".docx", ".txt"}
os.environ["PINECONE_API_KEY"] = settings.pinecone_api_key


@dataclass
class IngestionResult:
    filename: str
    file_type: str
    chunks_indexed: int
    pages_found: int
    status: str          
    error: str = ""


# ── Loaders: one function per file type ─────────────────────────

def load_pdf(path: str) -> List[Document]:
    """
    PyMuPDF gives us page number for free in metadata.
    Best for text-heavy PDFs, handles multi-column layout well.
    """
    path = str(Path(path).resolve())   
    loader = PyMuPDFLoader(path)
    return loader.load()


def load_docx(path: str) -> List[Document]:
    """
    python-docx gives us paragraphs and heading structure.
    We treat each paragraph as a separate document with para index.
    """
    path = str(Path(path).resolve())   
    doc = DocxDocument(path)
    documents = []

    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if not text:                    
            continue

        # detect if paragraph is a heading — useful metadata
        is_heading = para.style.name.startswith("Heading")
        heading_level = para.style.name if is_heading else None

        documents.append(Document(
            page_content=text,
            metadata={
                "paragraph_index": i,
                "is_heading": is_heading,
                "heading_level": heading_level,
                "page": 0,             # DOCX has no pages, normalise to 0
            }
        ))

    return documents


def load_txt(path: str) -> List[Document]:
    """
    TextLoader handles encoding detection.
    We split on double newlines to get logical paragraphs.
    """
    path = str(Path(path).resolve())
    # TextLoader returns a list with 1 Document by default
    loader = TextLoader(path, encoding="utf-8", autodetect_encoding=True)
    docs = loader.load()

    for doc in docs:
        doc.metadata["page"] = 0
        doc.metadata["total_pages"] = 1
        
    return docs


# ── Dispatcher: picks the right loader ──────────────────────────

def load_document(path: str) -> List[Document]:
    ext = Path(path).suffix.lower()

    if ext == ".pdf":
        return load_pdf(path)
    elif ext == ".docx":
        return load_docx(path)
    elif ext == ".txt":
        return load_txt(path)
    else:
        raise ValueError(f"Unsupported file type: {ext}")


# ── Metadata enrichment ──────────────────────────────────────────

def enrich_metadata(
    docs: List[Document],
    filename: str
) -> List[Document]:
    """
    Stamps every chunk with consistent metadata fields.
    These fields power citation tracking in Phase 3.
    """
    enriched = []
    for doc in docs:
        enriched.append(
            Document(
                page_content=doc.page_content,
                metadata={
                    "source": filename,             
                    "page": doc.metadata.get("page", 0), 
                }
            )
        )

    return enriched


# ── Validator: filters bad chunks before embedding ────────────────

def validate_docs(docs: List[Document], min_chars: int = 30) -> List[Document]:
    """
    Drop chunks that are too short to be meaningful.
    Avoids wasting Pinecone upserts on noise.
    """
    valid = []
    skipped = 0
    for doc in docs:
        content = doc.page_content.strip()
        if len(content) >= min_chars:
            doc.page_content = content   
            valid.append(doc)
        else:
            skipped += 1

    if skipped:
        print(f"Validator: dropped {skipped} chunks below {min_chars} chars")

    return valid


# ── Chunker ──────────────────────────────────────────────────────

def chunk_documents(docs: List[Document]) -> List[Document]:
    splitter = RecursiveCharacterTextSplitter(
        chunk_size=500,
        chunk_overlap=50,
        length_function=len,
        separators=["\n\n", "\n", ".", "!", "?", " "]
    )
    chunks = splitter.split_documents(docs)

    # stamp chunk index — needed for citation ordering later
    for i, chunk in enumerate(chunks):
        chunk.metadata["chunk_index"] = i

    return chunks


# ── Main ingestion function ──────────────────────────────────────

def ingest_file_from_path(filepath: str) -> IngestionResult:
    """
    Ingest a file that already exists on disk.
    Used by store_index.py for bulk ingestion.
    """
    filename = Path(filepath).name
    ext = Path(filepath).suffix.lower()
    file_type = ext.lstrip(".")

    if ext not in SUPPORTED_EXTENSIONS:
        return IngestionResult(
            filename=filename,
            file_type=file_type,
            chunks_indexed=0,
            pages_found=0,
            status="skipped",
            error=f"Unsupported extension: {ext}"
        )

    try:
        # Step 1 — Load
        raw_docs = load_document(filepath)
        pages_found = len(raw_docs)
        print(f"Loaded {pages_found} pages/blocks from {filename}")

        # Step 2 — Enrich metadata
        enriched = enrich_metadata(raw_docs, filename)

        # Step 3 — Validate
        valid = validate_docs(enriched)

        # Step 4 — Chunk
        chunks = chunk_documents(valid)
        print(f"Created {len(chunks)} chunks from {filename}")

        # Step 5 — Embed + upsert to Pinecone
        embeddings = HuggingFaceEmbeddings(model_name=settings.embed_model)
        
        index_name = settings.pinecone_index_name 
        pc = Pinecone(api_key=settings.pinecone_api_key)

        if not pc.has_index(index_name):
            pc.create_index(
                name=index_name,
                dimension=384,          # Dimension of embedding model
                metric="cosine",
                spec=ServerlessSpec(cloud="aws", region="us-east-1")
            )

        docsearch = PineconeVectorStore.from_documents(
            documents=chunks,
            embedding=embeddings,
            index_name=index_name,
            batch_size=100
        )

        print(f"Upserted {len(chunks)} chunks to Pinecone for {filename}")

        return IngestionResult(
            filename=filename,
            file_type=file_type,
            chunks_indexed=len(chunks),
            pages_found=pages_found,
            status="success"
        )

    except Exception as e:
        return IngestionResult(
            filename=filename,
            file_type=file_type,
            chunks_indexed=0,
            pages_found=0,
            status="failed",
            error=str(e)
        )


async def ingest_file_from_bytes(
    file_bytes: bytes,
    filename: str
) -> IngestionResult:
    """
    Ingest a file uploaded at runtime via FastAPI /ingest endpoint.
    Saves to a temp file, runs the same pipeline, then cleans up.
    """
    ext = Path(filename).suffix.lower()

    with tempfile.NamedTemporaryFile(suffix=ext, delete=False) as tmp:
        tmp.write(file_bytes)
        tmp_path = tmp.name

    try:
        return ingest_file_from_path(tmp_path)

    finally:
        os.unlink(tmp_path)     # always clean up regardless of success/fail